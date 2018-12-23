#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
import pandas as pd
import sys
import os
import string

#pip install python-docx
#pip install pandas
#pip install xlrd


def docx_replace(doc_obj, dikt):
    for p in doc_obj.paragraphs:
        p.text = string.Template(p.text).substitute(dikt)
                    
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace(cell, dikt)


def main(teplf,listf,outdir):
    xlsx = pd.read_excel(listf, skiprows=2, sheet_name="list").fillna('')
    ext = ".docx" 
    for i, row in xlsx.iterrows():
        doc = Document(teplf)
       
        name = row["name"]
        position = row["position"]
        
        docx_replace(doc, row.to_dict())
        print row

        name2 = name.replace(" ","")
        position2 = position.replace(" ","")
        
        outpath = u"{0}/{1}_{2}_{3}{4}".format(outdir,i,name2,position2,ext)
        doc.save(outpath)
        if os.path.exists(outpath):
            print(outpath)
    

    
if __name__ == "__main__":
    args = sys.argv
    args.pop(0)

    if len(args) < 3:
        print(u"引数が不足しています")
        
    teplf = args.pop(0)
    listf = args.pop(0)
    outdir = args.pop(0)
        
    main(teplf,listf,outdir)
